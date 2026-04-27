"""Generate updated Project Report with all benchmark changes."""
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# ═══ Title ═══
heading = doc.add_heading('Project Report: Performance Evaluation of Encryption Algorithms on Edge Devices', 0)
heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# ═══ 1. Problem Statement ═══
doc.add_heading('1. Problem Statement & Computing Strategy Used', level=1)
doc.add_heading('Problem Statement', level=2)
doc.add_paragraph(
    "IoT edge devices (e.g., smart sensors, medical monitors, RFID tags) are responsible for collecting highly sensitive data. "
    "To ensure privacy and data sovereignty, this information must be encrypted immediately at the edge before transmission. "
    "However, these edge devices operate under severe hardware constraints, possessing strictly limited CPU cycles, memory, and battery life. "
    "Traditional cryptographic standards (such as hardware-accelerated AES) may exhaust the resources of ultra-constrained devices. "
    "Thus, the core objective of this project is to conduct a rigorous comparative evaluation of standard versus lightweight encryption algorithms "
    "-- specifically measuring latency, throughput, CPU utilization, and memory footprint -- to identify the optimal algorithm for diverse edge environments."
)

doc.add_heading('Computing Strategy Used', level=2)
doc.add_paragraph(
    "The project employs an Edge-Heavy Distributed Processing strategy. "
    "Instead of centralizing the cryptographic workload in a powerful cloud server, encryption is offloaded directly to the simulated edge nodes. "
    "This strategy ensures data sovereignty and minimizes plaintext exposure over vulnerable networks. "
    "We utilized a Distributed Parallel Simulation driven by Python's ProcessPoolExecutor to simulate concurrent edge nodes operating in a geographic area. "
    "The network scaling is modeled using the fundamental Wireless Sensor Network (WSN) formula: N = A x D (Nodes = Area x Density). "
    "Finally, a Cloud/Fog Aggregator node receives the encrypted payloads, decrypts them, and verifies data integrity via SHA-256 hashing, while aggregating system-wide performance metrics."
)

doc.add_heading('Relevance to Distributed Systems', level=2)
p = doc.add_paragraph()
p.add_run("1. Distributed Edge Computing Architecture: ").bold = True
p.add_run("IoT edge devices form a distributed network where each node independently generates, encrypts, and transmits data, mirroring real-world distributed systems.\n")
p.add_run("2. Parallel Processing & Concurrency: ").bold = True
p.add_run("Python's ProcessPoolExecutor simulates N concurrent edge nodes, demonstrating process-level parallelism, workload distribution, and synchronization via futures.\n")
p.add_run("3. Communication & Data Integrity: ").bold = True
p.add_run("Encrypted payloads are transmitted from edge to cloud/fog aggregator, mimicking client-server message passing. HMAC/GCM authentication ensures data integrity.\n")
p.add_run("4. Scalability & Load Balancing: ").bold = True
p.add_run("Scalability experiments (5 to 1000 nodes) measure system throughput and latency variance. Low standard deviation confirms balanced workload distribution.\n")
p.add_run("5. Fault Tolerance & Data Sovereignty: ").bold = True
p.add_run("By encrypting at the edge (not centrally), data remains protected even if intermediate nodes are compromised.")

# ═══ 2. Literature Survey ═══
doc.add_heading('2. Detailed Literature Survey', level=1)
doc.add_paragraph("The following table summarizes the key research papers that inform and justify our project methodology:")
doc.add_paragraph()

survey_table = doc.add_table(rows=11, cols=4)
survey_table.style = 'Table Grid'
survey_table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["#", "Paper Title", "Authors / Year", "Key Contribution & Relevance to Our Project"]
for j, h in enumerate(headers):
    cell = survey_table.rows[0].cells[j]
    cell.text = h
    for p in cell.paragraphs:
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(10)

survey_data = [
    ["1", "The Design of Rijndael: AES", "Daemen & Rijmen, 2002",
     "Establishes the mathematical foundation of AES block cipher. We use AES-256-GCM as the industry-standard baseline, combining 14 rounds of SubBytes/ShiftRows/MixColumns with Galois/Counter Mode for authenticated encryption."],
    ["2", "ChaCha, a variant of Salsa20", "Bernstein D.J., 2008",
     "Introduces the ChaCha20 stream cipher using ARX (Add-Rotate-XOR) operations. Performs exceptionally well in software without requiring specialized hardware crypto instructions. Paired with Poly1305 MAC for authentication."],
    ["3", "PRESENT: An Ultra-Lightweight Block Cipher", "Bogdanov et al., CHES 2007",
     "Designs an SPN-based cipher for extreme constraints requiring only ~1,570 Gate Equivalents. Uses 31 rounds of 4-bit S-Box substitution + bit permutation. We implemented this from scratch to benchmark the lowest hardware tier."],
    ["4", "The SIMON and SPECK Families of Lightweight Block Ciphers", "Beaulieu et al., IACR 2013",
     "NSA-designed SPECK-64/128 ARX cipher achieves 24 cycles/byte on ARM Cortex. Uses 27 rounds of rotation + addition + XOR with no S-Boxes. We implemented this from scratch for software benchmarking."],
    ["5", "Lightweight Cryptography Algorithms for Resource-Constrained IoT Devices: A Review", "Thakor et al., IEEE Access 2021",
     "Provides a comprehensive comparative framework for benchmarking IoT encryption across latency, throughput, and memory. Our empirical measurement methodology directly mirrors the approach recommended in this review."],
    ["6", "Wireless Sensor Networks: A Survey", "Akyildiz et al., Computer Networks 2002",
     "Seminal paper establishing that sensor node count scales linearly with deployment area and density (N = A x D). This directly informs the distributed simulation scaling logic used in our project."],
    ["7", "Edge Computing: Vision and Challenges", "Shi et al., IEEE IoT Journal 2016",
     "Defines the edge computing paradigm where computation is pushed to the network edge. Justifies our architectural choice of performing encryption at the edge node rather than in the cloud, reducing latency and improving data sovereignty."],
    ["8", "The Internet of Things: A Survey", "Li et al., Information Systems Frontiers 2015",
     "Comprehensive survey of IoT architectures, protocols, and applications. Validates our synthetic sensor data model (temperature, humidity, pressure) and confirms the resource constraints we simulate for edge devices."],
    ["9", "Report on Lightweight Cryptography", "McKay et al., NIST IR 8114, 2017",
     "Official NIST report outlining criteria for selecting lightweight cryptographic algorithms for constrained devices. Frames our comparison criteria (throughput, memory, CPU) and validates the inclusion of PRESENT and SPECK as candidate algorithms."],
    ["10", "Review of Lightweight Block Ciphers", "Hatzivasilis et al., J. Cryptographic Engineering 2018",
     "Systematic review comparing lightweight block ciphers including PRESENT, SPECK, and SIMON across gate count, throughput, and security margins. Validates our from-scratch implementation approach and confirms the performance characteristics observed in our benchmarks."],
]

for i, row_data in enumerate(survey_data):
    for j, val in enumerate(row_data):
        cell = survey_table.rows[i+1].cells[j]
        cell.text = val
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

# ═══ 3. Implementation ═══
doc.add_heading('3. Implementation Details', level=1)

doc.add_heading('Tools and Technologies', level=2)
doc.add_paragraph(
    "The implementation was developed entirely in Python 3.10+. "
    "We utilized the PyCryptodome library for the industry-standard algorithms (AES-256-GCM and ChaCha20-Poly1305), "
    "while leveraging native Python constructs (Struct, Hashlib) to manually build the lightweight ciphers (PRESENT, SPECK) strictly according to their mathematical specifications. "
    "For performance profiling, we utilized high-resolution nanosecond timers (time.perf_counter()), psutil for exact CPU utilization, "
    "and tracemalloc for peak memory tracking. Data visualization was handled using Matplotlib and Pandas."
)

doc.add_heading('Dataset Generation', level=2)
doc.add_paragraph(
    "We generate synthetic IoT sensor data as JSON payloads containing timestamp, device_id, temperature_c, humidity_pct, pressure_hpa, light_lux, sensor_voltage, and status. "
    "Three data sizes were used: 1KB (~5 records, typical RFID payload), 10KB (~50 records, periodic sensor upload), and 100KB (~500 records, edge gateway aggregation). "
    "Records are accumulated until the serialized JSON meets the target byte size, then padded or trimmed to exactly match. "
    "The same data is used across all algorithms at each size for fair comparison."
)

doc.add_heading('Algorithm Methodology', level=2)
p = doc.add_paragraph()
p.add_run("AES-256-GCM: ").bold = True
p.add_run("Block cipher with 128-bit blocks, 256-bit key, 14 rounds of SubBytes/ShiftRows/MixColumns/AddRoundKey. Wrapped in Galois/Counter Mode for authenticated encryption. Implemented via PyCryptodome (uses hardware AES-NI).\n\n")
p.add_run("ChaCha20-Poly1305: ").bold = True
p.add_run("Stream cipher with 256-bit key, 20 rounds of quarter-round ARX operations generating a keystream XORed with plaintext. Poly1305 MAC provides authentication. Implemented via PyCryptodome.\n\n")
p.add_run("PRESENT-80 (Built from scratch): ").bold = True
p.add_run("Lightweight SPN block cipher with 64-bit blocks, 80-bit key, 31 rounds of 4-bit S-Box substitution + fixed P-Box bit permutation. Wrapped in CTR mode for arbitrary-length payloads. HMAC-SHA256 added for integrity.\n\n")
p.add_run("SPECK-64/128 (Built from scratch): ").bold = True
p.add_run("Lightweight ARX block cipher with 64-bit blocks, 128-bit key, 27 rounds using only bitwise rotation, modular addition, and XOR. No S-Boxes. Wrapped in CTR mode + HMAC-SHA256. Key expansion generates 27 round keys from master key.")

doc.add_heading('Benchmarking Pipeline', level=2)
doc.add_paragraph(
    "A custom BenchmarkProfiler context manager isolates the execution of each algorithm, capturing wall-clock time (time.perf_counter()), "
    "CPU utilization (psutil process CPU times), peak memory (tracemalloc), and throughput (bytes/second). "
    "A 50ms delay is injected between encrypt and decrypt to simulate real edge-to-cloud network latency. "
    "Each algorithm is tested 5 times per data size, and integrity is verified by asserting decrypted output matches original data. "
    "Total benchmark: 4 algorithms x 3 sizes x 5 iterations x 2 operations = 120 runs."
)

# ═══ 4. Results ═══
doc.add_heading('4. Results', level=1)

doc.add_heading('Encryption Performance Summary', level=2)
doc.add_paragraph("Averaged across all three data sizes (1KB, 10KB, 100KB) and 5 iterations:")

results_table = doc.add_table(rows=5, cols=6)
results_table.style = 'Table Grid'
results_table.alignment = WD_TABLE_ALIGNMENT.CENTER

res_headers = ["Algorithm", "Avg Time (ms)", "Avg CPU %", "Avg Memory (KB)", "Avg Throughput (MB/s)", "Composite Score"]
for j, h in enumerate(res_headers):
    cell = results_table.rows[0].cells[j]
    cell.text = h
    for p in cell.paragraphs:
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(10)

res_data = [
    ["AES-256-GCM", "1.06", "179.1", "76.5", "35.80", "6.0"],
    ["ChaCha20-Poly1305", "1.16", "143.1", "75.0", "22.69", "7.0"],
    ["PRESENT-80", "48,294.17", "88.6", "78.8", "0.001", "14.0"],
    ["SPECK-64/128", "2,091.74", "95.3", "78.7", "0.017", "13.0"],
]
for i, row_data in enumerate(res_data):
    for j, val in enumerate(row_data):
        cell = results_table.rows[i+1].cells[j]
        cell.text = val
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph("Composite Score: Each algorithm is ranked 1-4 on throughput (higher=better), latency (lower=better), CPU (lower=better), and memory (lower=better). The sum of ranks is the composite score. Lower score = better overall.")
doc.add_paragraph()

p = doc.add_paragraph()
p.add_run("Best Overall: AES-256-GCM (Score: 6.0). ").bold = True
p.add_run("AES wins on latency and throughput due to hardware AES-NI acceleration. ChaCha20 is the best choice on devices without AES-NI hardware acceleration.")

doc.add_heading('Key Observations from Graphs', level=2)
p = doc.add_paragraph()
p.add_run("Throughput: ").bold = True
p.add_run("AES achieves ~94 MB/s and ChaCha20 ~50 MB/s at 100KB. PRESENT and SPECK are orders of magnitude slower in pure Python.\n")
p.add_run("Encryption Time: ").bold = True
p.add_run("PRESENT takes ~127 seconds for 100KB. SPECK takes ~5.6 seconds. AES completes in ~1.2ms and ChaCha20 in ~2.1ms.\n")
p.add_run("CPU Usage: ").bold = True
p.add_run("All algorithms now show accurate measured CPU utilization (~88-179%) thanks to calibrated cpu_times() profiling with a calibration loop for sub-millisecond operations.\n")
p.add_run("Memory: ").bold = True
p.add_run("At 100KB: AES ~202KB, ChaCha20 ~201KB, PRESENT ~211KB, SPECK ~210KB peak memory.")

doc.add_heading('Why Single-Node Benchmarking AND Distributed Simulation?', level=2)
p2 = doc.add_paragraph()
p2.add_run("Single-Node Benchmarking: ").bold = True
p2.add_run("Establishes the performance baseline for each algorithm in isolation. Without this, distributed scalability cannot be measured. "
    "Testing across 1KB, 10KB, and 100KB payloads reveals each algorithm's setup overhead vs. processing cost curve, identifying the ideal network packet size. "
    "It also isolates exact CPU% and peak RAM per device, proving whether the algorithm is physically deployable on constrained IoT hardware.\n\n")
p2.add_run("Distributed Simulation: ").bold = True
p2.add_run("Real IoT deployments involve thousands of nodes encrypting concurrently. A single-node test cannot capture inter-node overhead or parallel contention. "
    "By distributing a fixed 100KB workload across N=5 to N=1000 nodes, we prove near-linear throughput scaling for AES/ChaCha20. "
    "Key finding: AES and ChaCha20 scale almost perfectly, while PRESENT and SPECK's process overhead exceeds their encryption time at high node counts.")

# ═══ 5. Conclusion ═══
doc.add_heading('5. Conclusion', level=1)

doc.add_heading('Where Each Algorithm Excels and Fails', level=2)

conc_table = doc.add_table(rows=5, cols=4)
conc_table.style = 'Table Grid'
conc_table.alignment = WD_TABLE_ALIGNMENT.CENTER

conc_headers = ["Algorithm", "Best For", "Fails When", "Why"]
for j, h in enumerate(conc_headers):
    cell = conc_table.rows[0].cells[j]
    cell.text = h
    for p in cell.paragraphs:
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(10)

conc_data = [
    ["AES-256-GCM", "High-performance edge nodes (Raspberry Pi, x86 with AES-NI)", "Ultra-constrained devices (passive RFID, 8-bit MCUs)", "Requires ~50,000 GE and dedicated hardware instructions for speed"],
    ["ChaCha20-Poly1305", "Software-only ARM/RISC-V MCUs (no crypto hardware needed)", "Extreme gate-count limits (needs ~38,000 GE)", "Pure ARX ops are fast in software but still require moderate silicon area"],
    ["PRESENT-80", "Passive RFID, smart cards (needs only ~1,570 GE)", "Any payload >1KB in pure software (~375ms per 1KB)", "Designed for hardware ASIC, not software. 31 rounds of S-Box are slow interpreted."],
    ["SPECK-64/128", "Constrained MCUs needing speed (24 cycles/byte on ARM)", "Large payloads in pure Python (~1.7s per 100KB)", "Optimized for C/ASM on MCUs. Python interpreter overhead negates ARX advantages."],
]
for i, row_data in enumerate(conc_data):
    for j, val in enumerate(row_data):
        cell = conc_table.rows[i+1].cells[j]
        cell.text = val
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

doc.add_paragraph()
doc.add_paragraph(
    "In conclusion, there is no single 'best' encryption algorithm for all edge devices. The optimal choice depends critically on the hardware capabilities "
    "of the target device. For devices with hardware crypto acceleration (AES-NI), AES-256-GCM is unmatched. For software-only microcontrollers, "
    "ChaCha20-Poly1305 provides the best throughput. For ultra-constrained hardware (RFID), PRESENT achieves security with minimal gate count. "
    "And for microcontrollers requiring the best software speed-to-gate-count ratio, SPECK is the optimal choice."
)

# ═══ 6. References ═══
doc.add_heading('6. References', level=1)
refs = [
    "[1] Daemen, J. & Rijmen, V. (2002). \"The Design of Rijndael: AES.\" Springer.",
    "[2] Bernstein, D.J. (2008). \"ChaCha, a variant of Salsa20.\" SASC 2008.",
    "[3] Bogdanov, A. et al. (2007). \"PRESENT: An Ultra-Lightweight Block Cipher.\" CHES 2007, LNCS 4727.",
    "[4] Beaulieu, R. et al. (2013). \"The SIMON and SPECK Families of Lightweight Block Ciphers.\" IACR ePrint 2013/404.",
    "[5] Thakor, V. et al. (2021). \"Lightweight Cryptography Algorithms for Resource-Constrained IoT Devices: A Review.\" IEEE Access, Vol. 9.",
    "[6] Akyildiz, I.F. et al. (2002). \"Wireless Sensor Networks: A Survey.\" Computer Networks, Vol. 38(4), pp. 393-422.",
    "[7] Shi, W. et al. (2016). \"Edge Computing: Vision and Challenges.\" IEEE Internet of Things Journal, 3(5), pp. 637-646.",
    "[8] Li, S. et al. (2015). \"The Internet of Things: A Survey.\" Information Systems Frontiers, 17(2), pp. 243-259.",
    "[9] McKay, K. et al. (2017). \"Report on Lightweight Cryptography.\" NIST Internal Report 8114.",
    "[10] Hatzivasilis, G. et al. (2018). \"Review of Lightweight Block Ciphers.\" Journal of Cryptographic Engineering, 8(2), pp. 141-184.",
    "",
    "Online References:",
    "[W1] https://en.wikipedia.org/wiki/Advanced_Encryption_Standard",
    "[W2] https://en.wikipedia.org/wiki/Salsa20#ChaCha_variant",
    "[W3] https://en.wikipedia.org/wiki/PRESENT_(cipher)",
]
for ref in refs:
    doc.add_paragraph(ref)

doc.add_paragraph()
doc.add_paragraph("Date: April 2026")

out_path = os.path.join(os.path.dirname(__file__), "Project_Report.docx")
doc.save(out_path)
print(f"Report saved to: {out_path}")
